import streamlit as st
import pandas as pd
import sqlite3
from io import BytesIO
from datetime import datetime

# Word
from docx import Document

# PDF عربي
import arabic_reshaper
from bidi.algorithm import get_display

from reportlab.platypus import (
    SimpleDocTemplate,
    Table,
    TableStyle,
    Paragraph,
    Spacer
)

from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4

from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Excel
from openpyxl import Workbook

# =========================================================
# إعدادات الصفحة
# =========================================================
st.set_page_config(
    page_title="نظام جماعة معلين",
    page_icon="⚖️",
    layout="wide"
)

# =========================================================
# التصميم الاحترافي
# =========================================================
st.markdown("""
<style>

html, body, [data-testid="stAppViewContainer"], .main {
    direction: rtl;
    text-align: right;
    background-color: #f5f7fb;
}

/* عناوين */
h1,h2,h3,h4,h5,h6{
    font-weight:bold;
    color:#0f172a;
}

/* بطاقات الإحصائيات */
.stat-card{
    background: linear-gradient(135deg,#1e3c72,#2a5298);
    padding:25px;
    border-radius:20px;
    color:white;
    text-align:center;
    box-shadow:0px 5px 20px rgba(0,0,0,0.15);
    transition:0.3s;
    margin-bottom:15px;
}

.stat-card:hover{
    transform:scale(1.03);
}

/* الأزرار */
.stButton>button{
    width:100%;
    border:none;
    border-radius:12px;
    background:linear-gradient(90deg,#1565c0,#1976d2);
    color:white;
    font-size:16px;
    font-weight:bold;
    padding:10px;
}

/* المدخلات */
.stTextInput input,
.stNumberInput input{
    border-radius:10px;
}

/* شاشة الدخول */
.login-box{
    background:white;
    padding:40px;
    border-radius:20px;
    box-shadow:0px 0px 25px rgba(0,0,0,0.1);
    margin-top:80px;
}

</style>
""", unsafe_allow_html=True)

# =========================================================
# قاعدة البيانات
# =========================================================
conn = sqlite3.connect(
    "maalin.db",
    check_same_thread=False
)

cursor = conn.cursor()

# =========================================================
# إنشاء الجداول
# =========================================================
cursor.execute("""
CREATE TABLE IF NOT EXISTS members(
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    name TEXT,
    family_code TEXT,
    gender TEXT,
    paid TEXT,
    phone TEXT,
    notes TEXT,
    created_at TEXT
)
""")

cursor.execute("""
CREATE TABLE IF NOT EXISTS users(
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    username TEXT,
    password TEXT
)
""")

conn.commit()

# =========================================================
# إنشاء الأدمن
# =========================================================
cursor.execute("""
SELECT * FROM users
WHERE username='admin'
""")

admin_exists = cursor.fetchone()

if not admin_exists:

    cursor.execute("""
    INSERT INTO users(username,password)
    VALUES('admin','123')
    """)

    conn.commit()

# =========================================================
# دوال
# =========================================================
def load_members():

    return pd.read_sql(
        "SELECT * FROM members",
        conn
    )


# =========================================================
# Excel
# =========================================================
def create_excel(df):

    output = BytesIO()

    wb = Workbook()

    ws = wb.active

    ws.title = "الأعضاء"

    ws.append(list(df.columns))

    for row in df.values.tolist():
        ws.append(row)

    wb.save(output)

    return output.getvalue()


# =========================================================
# Word
# =========================================================
def create_word(df):

    doc = Document()

    doc.add_heading(
        "تقرير أعضاء جماعة معلين",
        level=1
    )

    table = doc.add_table(
        rows=1,
        cols=len(df.columns)
    )

    table.style = "Table Grid"

    hdr = table.rows[0].cells

    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)

    for _, row in df.iterrows():

        cells = table.add_row().cells

        for i, val in enumerate(row):
            cells[i].text = str(val)

    output = BytesIO()

    doc.save(output)

    return output.getvalue()


# =========================================================
# PDF عربي
# =========================================================
def create_pdf(df):

    output = BytesIO()

    pdfmetrics.registerFont(
        TTFont(
            'Arabic',
            'Amiri-Regular.ttf'
        )
    )

    doc = SimpleDocTemplate(
        output,
        pagesize=A4
    )

    styles = getSampleStyleSheet()

    elements = []

    title_text = "تقرير أعضاء جماعة معلين"

    reshaped_title = arabic_reshaper.reshape(
        title_text
    )

    bidi_title = get_display(
        reshaped_title
    )

    title = Paragraph(
        f"<font name='Arabic'>{bidi_title}</font>",
        styles['Title']
    )

    elements.append(title)

    elements.append(Spacer(1,20))

    data = []

    headers = []

    for col in df.columns:

        reshaped = arabic_reshaper.reshape(
            str(col)
        )

        bidi_text = get_display(
            reshaped
        )

        headers.append(bidi_text)

    data.append(headers)

    for _, row in df.iterrows():

        row_data = []

        for item in row:

            reshaped = arabic_reshaper.reshape(
                str(item)
            )

            bidi_text = get_display(
                reshaped
            )

            row_data.append(bidi_text)

        data.append(row_data)

    table = Table(data)

    table.setStyle(TableStyle([
        ('BACKGROUND',(0,0),(-1,0),colors.darkblue),
        ('TEXTCOLOR',(0,0),(-1,0),colors.white),
        ('GRID',(0,0),(-1,-1),1,colors.black),
        ('FONTNAME',(0,0),(-1,-1),'Arabic'),
        ('FONTSIZE',(0,0),(-1,-1),12),
        ('ALIGN',(0,0),(-1,-1),'RIGHT'),
        ('BACKGROUND',(0,1),(-1,-1),colors.beige),
    ]))

    elements.append(table)

    doc.build(elements)

    return output.getvalue()

# =========================================================
# session
# =========================================================
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False

# =========================================================
# تسجيل الدخول
# =========================================================
if not st.session_state.logged_in:

    col1, col2, col3 = st.columns([1,2,1])

    with col2:

        st.markdown("""
        <div class="login-box">
        <h1 style="text-align:center">
        ⚖️ نظام جماعة معلين
        </h1>
        <p style="text-align:center;color:gray">
        تسجيل الدخول للإدارة
        </p>
        </div>
        """, unsafe_allow_html=True)

        username = st.text_input(
            "👤 اسم المستخدم"
        )

        password = st.text_input(
            "🔑 كلمة المرور",
            type="password"
        )

        if st.button("🚀 دخول النظام"):

            cursor.execute("""
            SELECT * FROM users
            WHERE username=? AND password=?
            """, (username, password))

            user = cursor.fetchone()

            if user:

                st.session_state.logged_in = True

                st.rerun()

            else:
                st.error("بيانات الدخول غير صحيحة")

# =========================================================
# النظام الرئيسي
# =========================================================
else:

    col1, col2 = st.columns([8,1])

    with col2:

        if st.button("🚪 خروج"):

            st.session_state.logged_in = False

            st.rerun()

    st.title("⚖️ نظام جماعة معلين الاحترافي")

    df = load_members()

    total = len(df)

    paid = len(df[df["paid"] == "نعم"])

    not_paid = total - paid

    males = len(df[df["gender"] == "ذكر"])

    females = len(df[df["gender"] == "أنثى"])

    fund = paid * 500

    st.markdown("## 📊 الإحصائيات")

    c1, c2, c3, c4, c5, c6 = st.columns(6)

    c1.markdown(f"""
    <div class='stat-card'>
    👥
    <h2>{total}</h2>
    إجمالي الأعضاء
    </div>
    """, unsafe_allow_html=True)

    c2.markdown(f"""
    <div class='stat-card'>
    👨
    <h2>{males}</h2>
    الذكور
    </div>
    """, unsafe_allow_html=True)

    c3.markdown(f"""
    <div class='stat-card'>
    👩
    <h2>{females}</h2>
    الإناث
    </div>
    """, unsafe_allow_html=True)

    c4.markdown(f"""
    <div class='stat-card'>
    ✅
    <h2>{paid}</h2>
    المسددين
    </div>
    """, unsafe_allow_html=True)

    c5.markdown(f"""
    <div class='stat-card'>
    ❌
    <h2>{not_paid}</h2>
    غير المسددين
    </div>
    """, unsafe_allow_html=True)

    c6.markdown(f"""
    <div class='stat-card'>
    💰
    <h2>{fund:,}</h2>
    ريال
    </div>
    """, unsafe_allow_html=True)

    st.markdown("---")

    tab1, tab2, tab3 = st.tabs([
        "👥 إدارة الأعضاء",
        "📊 التقارير",
        "💰 التقسيم المالي"
    ])

    # =====================================================
    # الأعضاء
    # =====================================================
    with tab1:

        st.subheader("➕ إضافة عضو")

        with st.form("add_member"):

            name = st.text_input("الاسم")

            family = st.text_input("كود العائلة")

            gender = st.selectbox(
                "الجنس",
                ["ذكر","أنثى"]
            )

            paid = st.selectbox(
                "حالة الدفع",
                ["نعم","لا"]
            )

            phone = st.text_input("الجوال")

            notes = st.text_area("ملاحظات")

            submit = st.form_submit_button(
                "إضافة عضو"
            )

            if submit:

                cursor.execute("""
                INSERT INTO members(
                    name,
                    family_code,
                    gender,
                    paid,
                    phone,
                    notes,
                    created_at
                )
                VALUES(?,?,?,?,?,?,?)
                """, (
                    name,
                    family,
                    gender,
                    paid,
                    phone,
                    notes,
                    datetime.now().strftime("%Y-%m-%d")
                ))

                conn.commit()

                st.success("تمت إضافة العضو")

                st.rerun()

        st.markdown("---")

        st.subheader("✏️ تعديل وحذف الأعضاء")

        members = load_members()

        for _, row in members.iterrows():

            with st.expander(f"👤 {row['name']}"):

                new_name = st.text_input(
                    "الاسم",
                    row["name"],
                    key=f"name_{row['id']}"
                )

                new_family = st.text_input(
                    "العائلة",
                    row["family_code"],
                    key=f"family_{row['id']}"
                )

                new_gender = st.selectbox(
                    "الجنس",
                    ["ذكر","أنثى"],
                    index=0 if row["gender"]=="ذكر" else 1,
                    key=f"gender_{row['id']}"
                )

                new_paid = st.selectbox(
                    "الدفع",
                    ["نعم","لا"],
                    index=0 if row["paid"]=="نعم" else 1,
                    key=f"paid_{row['id']}"
                )

                new_phone = st.text_input(
                    "الجوال",
                    row["phone"],
                    key=f"phone_{row['id']}"
                )

                new_notes = st.text_area(
                    "ملاحظات",
                    row["notes"],
                    key=f"notes_{row['id']}"
                )

                cc1, cc2 = st.columns(2)

                with cc1:

                    if st.button(
                        "💾 حفظ",
                        key=f"save_{row['id']}"
                    ):

                        cursor.execute("""
                        UPDATE members
                        SET
                        name=?,
                        family_code=?,
                        gender=?,
                        paid=?,
                        phone=?,
                        notes=?
                        WHERE id=?
                        """, (
                            new_name,
                            new_family,
                            new_gender,
                            new_paid,
                            new_phone,
                            new_notes,
                            row["id"]
                        ))

                        conn.commit()

                        st.success("تم التعديل")

                        st.rerun()

                with cc2:

                    if st.button(
                        "🗑️ حذف",
                        key=f"del_{row['id']}"
                    ):

                        cursor.execute("""
                        DELETE FROM members
                        WHERE id=?
                        """, (row["id"],))

                        conn.commit()

                        st.warning("تم حذف العضو")

                        st.rerun()

    # =====================================================
    # التقارير
    # =====================================================
    with tab2:

        st.subheader("📊 التقارير")

        report_df = load_members()

        st.dataframe(
            report_df,
            use_container_width=True
        )

        # Excel
        excel_file = create_excel(report_df)

        st.download_button(
            "📥 تحميل Excel",
            excel_file,
            "report.xlsx"
        )

        # Word
        word_file = create_word(report_df)

        st.download_button(
            "📥 تحميل Word",
            word_file,
            "report.docx"
        )

        # PDF
        pdf_file = create_pdf(report_df)

        st.download_button(
            "📥 تحميل PDF",
            pdf_file,
            "report.pdf"
        )

    # =====================================================
    # التقسيم المالي
    # =====================================================
    with tab3:

        st.subheader("💰 التقسيم المالي")

        amount = st.number_input(
            "المبلغ الإجمالي",
            min_value=0.0
        )

        if amount > 0 and total > 0:

            share = amount / total

            st.metric(
                "المبلغ للفرد الواحد",
                f"{share:,.2f} ريال"
            )

            calc_df = load_members()

            calc_df["المبلغ المستحق"] = round(
                share,
                2
            )

            st.dataframe(
                calc_df,
                use_container_width=True
            )
